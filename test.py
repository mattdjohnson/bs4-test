html_doc = """<body><div class="Section">
<p class="MsoTitle"><strong>SYSTEM SECURITY PLAN (SSP)</strong></p>
<p class="MsoSubtitle"><strong>National Reconnaissance Office</strong></p>
<p class="MsoNormal">Hey dude</p>
<p class="MsoNormal"><span style="font-size: 12pt;">Project ID: <strong>200410</strong><span></p>
<p class="MsoNormal"><span style="font-size: 12pt;">Project Name: <strong id="project_id">UECM</strong></span></p>
<p class="MsoNormal"><span style="font-size: 12pt;">Date: <strong>08/20/2021</strong></span></p>
<span style="text-decoration: underline;">SOME TEXT In just a span tag</span>
<p class="MsoNormal">Normal not underlined <span style="text-decoration: underline;"> underline in msoNormal</span></p>
</div>
</body>
"""

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()

soup = BeautifulSoup(html_doc, 'html.parser')


def process_style (tag, docObj):
    style = tag.get("style")
    # print(f"INSIDE process_style:{style} TYPE:{type(style)} TAG:{tag} STRINGS:{list(tag.parent.strings)}")
    # font-size will be handled as WD_STYLE_TYPE.CHARACTER, like MsoNormal

    
    if style:
        if list(style.split(":"))[0] == 'text-decoration':
            if list(style.split(": "))[1] == 'underline;':
                # DRAGONS!!
                # docObj.add_text(tag.string)
                # docObj.add_run('from inside process_style').underline = True
                

    # <p style="text-align: left;">SOME TEXT TO TEST</p>
    # <p style="text-align: center;">SOME TEXT TO TEST</p>
    # <p style="text-align: right;">SOME TEXT TO TEST</p>
    # <p style="padding-left: 40px;">SOME TEXT TO TEST</p>
    # DOCX: paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER https://python-docx.readthedocs.io/en/latest/api/enum/WdAlignParagraph.html
    pass

def process_inline_tags (tag, docObj):
    if tag.children:
        for sub_tag in tag.children:
            if sub_tag.name == 'span':
                # <span style="text-decoration: underline;">SOME TEXT</span>
                # STRIKETHROUGH <span style="text-decoration: line-through;">completion</span>
                # <span style="color: #993300;">SOME COLORED TEXT</span>
                # <span style="background-color: #ffff00;">SOME TEXT TO TEST</span>

                process_style(sub_tag, docObj)

                # print(f"INSIDE SPAN:{sub_tag}")
                if sub_tag.children:
                    process_inline_tags(sub_tag, docObj)
            elif sub_tag.name == 'em':
                # print(f"INSIDE EM:{sub_tag}")
                if sub_tag.children:
                    process_inline_tags(sub_tag, docObj)
            elif sub_tag.name == 'strong':
                # print(f"INSIDE STRONG:{sub_tag}")
                if sub_tag.children:
                    process_inline_tags(sub_tag, docObj)



def process_ul (tag):
    """
        <ul style="list-style-position: inside;">
            <li>ONE</li>
            <li>TWO</li>
        </ul>
    """
    pass

def process_ol (tag):
    """
    <ol style="list-style-position: inside;">
        <li>&nbsp;FIRST ITEM<br /></li>
        <li>&nbsp;SECOND ITEM</li>
    </ol>
    """
    pass

# PROCESS TABLES

"""
    <blockquote>
        <p>A blockquote to process</p>
    </blockquote>

    <p><a title="A HYPERLINK!" href="#">A HYPERLINK!</a><br /></p>
"""


for tag in soup.find_all(True):
    # print(f"{len(list(tag.children))} TAG:{tag.name} CONTENTS:{tag.contents} ATTRS:{tag.attrs}")
    # check for Section Div
    if tag.name == 'body':
        continue
    if tag.name == 'div':
        if tag.attrs:
            #print(f"{tag} ATTRS:{tag.attrs}")
            for class_val in tag.get("class"):
                if class_val == 'Section':
                    # Process Word Doc Section, no action now https://python-docx.readthedocs.io/en/latest/user/sections.html?highlight=document.add_section#adding-a-new-section
                    pass
    if tag.name == 'p':
        if tag.attrs:
            for class_val in tag.get("class"):
                if class_val == 'MsoTitle':
                    p = document.add_paragraph(tag.string, style='Title')                   
                elif class_val == 'MsoSubtitle':       
                    p = document.add_paragraph(tag.string, style='Subtitle')
                elif class_val == 'MsoNormal':
                    p = document.add_paragraph(tag.string, style='Normal')
                    if tag.children:
                        process_inline_tags (tag, p)
                        # for sub_tag in tag.children:
                        #     print(f"SUB_TAG:{sub_tag} TYPE:{type(sub_tag)} NAME:{sub_tag.name}")

# <p class="MsoNormal"><span style="font-size: 12pt;">Date: <strong>08/20/2021</strong></span></p>

# https://stackoverflow.com/questions/27884703/set-paragraph-font-in-python-docx
font_styles = document.styles
font_charstyle = font_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
font_object = font_charstyle.font
font_object.size = Pt(20)
font_object.name = 'Times New Roman'

p = document.add_paragraph()
p.add_run('Date: ', style='CommentsStyle')
p.add_run('08/20/2021').bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic ').italic = True
p.add_run(' and some ')
p.add_run('underline this!').underline = True

document.save('demo.docx')